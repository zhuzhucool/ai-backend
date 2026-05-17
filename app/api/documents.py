from fastapi import APIRouter, Depends, HTTPException, UploadFile, Query, File
from app.schemas.document import  DocumentUploadResponse
from app.models.chat_message import ChatMessage
from app.models.document import Document
from app.db.session import get_session
from sqlmodel import Session, select
from app.core.security import verify_api_key, get_current_user_id
from app.core import config
import time
from pathlib import Path
from io import BytesIO
from pypdf import PdfReader
from docx import Document as DocxDocument
import openpyxl
import uuid
from sqlalchemy import func
import tempfile, os
from app.parsers import get_parser

settings = config.Settings()
router = APIRouter()
UPLOAD_DIR = Path("uploads")
ALLOWED_EXTENSIONS = (".txt", ".pdf", ".docx", ".xlsx")
MAX_SIZE = 20 * 1024 * 1024  # 20MB


def extract_text_from_bytes(content: bytes, ext: str) -> str:
    if ext == ".txt":
        return content.decode("utf-8")

    elif ext == ".pdf":
        reader = PdfReader(BytesIO(content))
        return "\n".join(page.extract_text() or "" for page in reader.pages)

    elif ext == ".docx":
        doc = DocxDocument(BytesIO(content))
        return "\n".join(para.text for para in doc.paragraphs)

    elif ext == ".xlsx":
        wb = openpyxl.load_workbook(BytesIO(content), read_only=True)
        rows = []
        for sheet in wb.sheetnames:
            for row in wb[sheet].iter_rows(values_only=True):
                cells = [str(c) for c in row if c is not None]
                if cells:
                    rows.append(" | ".join(cells))
        return "\n".join(rows)

    return ""


# @router.post("/documents/upload", response_model=DocumentUploadResponse, dependencies=[Depends(verify_api_key)])
# async def upload_document(
#     file: UploadFile,
#     db: Session = Depends(get_session),
#     user_id: str = Depends(get_current_user_id)
#     ):

#     # todo 优化pdf解析速度

#     status = "uploaded"
#     start = time.perf_counter()
#     # 1. 校验文件名
#     if not file.filename:
#         raise HTTPException(status_code=400, detail="文件名不能为空")

#     # 2. 校验格式
#     if not file.filename.lower().endswith(ALLOWED_EXTENSIONS):
#         raise HTTPException(status_code=400, detail="仅支持 txt/pdf/docx/xlsx 格式")

#     # 3. 读取内容并校验大小（用实际读取的字节判断，file.size 不一定可靠）
#     read_start = time.perf_counter()
#     content = await file.read()
#     read_ms = int((time.perf_counter() - read_start) * 1000)
#     if len(content) > MAX_SIZE:
#         raise HTTPException(status_code=400, detail=f"文件大小超过限制（最大 {MAX_SIZE // 1024 // 1024}MB）")

#     # 4. 提取文本内容
#     parse_start = time.perf_counter()
#     text_content = None
#     ext = Path(file.filename).suffix.lower()
#     if ext in (".txt", ".pdf", ".docx", ".xlsx"):
#         try:
#             text_content = extract_text_from_bytes(content, ext)
#         except Exception:
#             text_content = None  # 提取失败不阻断上传
#             status = "failed"
#     parse_ms = int((time.perf_counter() - parse_start) * 1000)


#     UPLOAD_DIR.mkdir(exist_ok=True) #若路径不存在则创建
#     safe_filename = f"{uuid.uuid4().hex}_{Path(file.filename).name}" #避免文件名重复
#     file_path = UPLOAD_DIR / safe_filename #避免文件名重复
    
#     #复制文件内容
#     write_start = time.perf_counter()
#     with file_path.open("wb") as buffer:
#         buffer.write(content)
#     write_ms = int((time.perf_counter() - write_start) * 1000)

#     document = Document(
#             user_id=user_id,
#             filename=file.filename,
#             file_path=str(file_path),
#             content_type=file.content_type,
#             size=len(content),
#             text_content=text_content,
#             status=status
#         )
#     #上传数据库并返回id
#     db_start = time.perf_counter()
#     db.add(document)
#     db.commit()
#     db.refresh(document)
#     db_ms = int((time.perf_counter() - db_start) * 1000)
#     total_ms = int((time.perf_counter() - start) * 1000)
#     print(f"upload_document read={read_ms}ms parse={parse_ms}ms write={write_ms}ms db={db_ms}ms total={total_ms}ms file={file.filename}")

#     return DocumentUploadResponse(
#         document_id=document.id,
#         filename=document.filename,
#         status=document.status
#     )


@router.post("/documents/upload")
async def upload_document(file: UploadFile = File(...)):
    # 1. 保存文件
    with tempfile.NamedTemporaryFile(delete=False, suffix=file.filename) as tmp:
        content = await file.read()
        tmp.write(content)
        tmp_path = tmp.name
    
    # 2. 解析
    parser = get_parser(file.filename)
    chunks = parser.parse(tmp_path, file.filename)
    
    # 3. 清理临时文件
    os.unlink(tmp_path)
    
    # 4. 返回结果（后续步骤会改成入库）
    return {
        "filename": file.filename,
        "chunks_count": len(chunks),
        "sample": chunks[0].text[:200] if chunks else "",
        "chunks": chunks
    }

@router.get("/documents", response_model=dict)
async def list_documents(
    user_id: int,
    db: Session = Depends(get_session),
    limit: int = Query(default=20, ge=1, le=100),
    offset: int = Query(default=0, ge=0)
):
    if user_id == None:
        return {"user_id": user_id, "error" : "当前user_id不存在"}
    # 查总数
    stmt = (
        select(func.count(Document.id))
        .where(Document.user_id == user_id)
    )
    total = db.exec(stmt).one()

    stmt = (
        select(ChatMessage)
        .where(Document.user_id == user_id)
        .order_by(Document.created_at.desc())
        .offset(offset)
        .limit(limit)
    )
    documents = db.exec(stmt).all()

    return {
        "total": total,
        "documents": documents
    }