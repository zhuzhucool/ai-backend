from fastapi import Depends, FastAPI, HTTPException
from pydantic import BaseModel, Field
from app.core import config
from app.services import llm
from contextlib import asynccontextmanager
from app.db.session import check_database_connection
from app.db.init_db import init_db
from app.db.session import get_session
from sqlmodel import Session, select
from app.models.chat_message import ChatMessage
from app.models.llm_log import LLMLog
from fastapi import FastAPI, File, UploadFile
import time
from pathlib import Path
import shutil
import uuid
from app.models.document import Document
from io import BytesIO
from pypdf import PdfReader
from docx import Document as DocxDocument
import openpyxl
from datetime import datetime
from fastapi import Query
from sqlalchemy import func



@asynccontextmanager
async def lifespan(app: FastAPI):
    check_database_connection()
    init_db()
    yield


app = FastAPI(lifespan=lifespan)

settings = config.Settings()
UPLOAD_DIR = Path("uploads")
ALLOWED_EXTENSIONS = (".txt", ".pdf", ".docx", ".xlsx")

MAX_SIZE = 20 * 1024 * 1024  # 20MB

print(settings)
# 定义请求体数据模型
class Item(BaseModel):
    name: str
    description: str | None = None
    price: float
    tax: float | None = None
class ChatRequest(BaseModel):
    user_id: int
    message: str 
    session_id: int = "123"
    temperature: float = Field(ge=0.0, le=2.0)
    max_tokens: int = Field(ge=1, le=1024*3)

class Usage(BaseModel):
    prompt_tokens: int = 0
    completion_tokens: int = 0
    total_tokens: int = 0

class ChatResponse(BaseModel):
    session_id: int
    message: str
    model: str
    usage: Usage | None

class DocumentUploadResponse(BaseModel):
    document_id: int
    filename: str
    status: str

class DocumentResponse(BaseModel):
    id: int
    filename: str
    status: str
    created_at: datetime


def extract_text_from_bytes(content: bytes, ext: str) -> str:
    if ext == ".txt":
        return content.decode("utf-8")

    elif ext == ".pdf":
        reader = PdfReader(BytesIO(content))
        return "\n".join(page.extract_text() or "" for page in reader.pages)

    elif ext == ".docx":
        doc = DocxDocument(BytesIO(content))
        return "\n".join(para.text for para in doc.paragraphs)

    elif ext == ".xlsx":
        wb = openpyxl.load_workbook(BytesIO(content), read_only=True)
        rows = []
        for sheet in wb.sheetnames:
            for row in wb[sheet].iter_rows(values_only=True):
                cells = [str(c) for c in row if c is not None]
                if cells:
                    rows.append(" | ".join(cells))
        return "\n".join(rows)

    return ""
    


@app.get("/documents", response_model=dict)
async def list_documents(
    user_id: int,
    db: Session = Depends(get_session),
    limit: int = Query(default=20, ge=1, le=100),
    offset: int = Query(default=0, ge=0)
):
    if user_id == None:
        return {"user_id": user_id, "error" : "当前user_id不存在"}
    # 查总数
    stmt = (
        select(func.count(Document.id))
        .where(Document.user_id == user_id)
    )
    total = db.exec(stmt).one()

    stmt = (
        select(ChatMessage)
        .where(Document.user_id == user_id)
        .order_by(Document.created_at.desc())
        .offset(offset)
        .limit(limit)
    )
    documents = db.exec(stmt).all()

    return {
        "total": total,
        "documents": documents
    }


@app.get("/health")
async def chat_health():
    return {"status": "ok", "service" : "ai-backend"}



@app.post("/chat", response_model=ChatResponse)
async def chat(req: ChatRequest, db: Session = Depends(get_session)):
    # TODO: 后面补事务、失败日志、Redis 限流/缓存、异步任务队列
    if not req.message.strip():
        raise HTTPException(status_code=400, detail={
            "error": "bad_request",
            "message": "message 不能为空"
        })
    
    user_message = ChatMessage(
        session_id=req.session_id,
        user_id=req.user_id,
        content=req.message,
        role="user"
    )

    db.add(user_message) #准备新增
    db.commit() #真正写入数据库
    db.refresh(user_message) #把数据库生成的ID信息读回来

    #日志记录
    start = time.perf_counter()
    success = False
    error_message = None
    response = None

    try:

        response = llm.llm_chat(req.message, req.temperature, req.max_tokens)
        success = True
    except llm.LLMError as e:
        error_message = e.message
        raise HTTPException(status_code=e.status_code, detail=e.message) from e
    finally:
        latency_ms = int((time.perf_counter() - start) * 1000)

    # 组装日志
    log = LLMLog(
        user_id=req.user_id,
        session_id=req.session_id,
        model=settings.OPENAI_MODEL,
        prompt_tokens=response.usage.prompt_tokens if response and response.usage else 0,
        completion_tokens=response.usage.completion_tokens if response and response.usage else 0,
        total_tokens=response.usage.total_tokens if response and response.usage else 0,
        latency_ms=latency_ms,
        success=success,
        error_message=error_message,
    )
    # 写入
    db.add(log)
    db.commit()

    assistant_content = response.choices[0].message.content
    assistant_message = ChatMessage(
        session_id=req.session_id,
        user_id=req.user_id,
        content=assistant_content,
        role="assistant"
    )
    db.add(assistant_message) #准备新增
    db.commit() #真正写入数据库
    db.refresh(assistant_message) #把数据库生成的ID信息读回来

    result = ChatResponse(session_id=req.session_id,
                          message=response.choices[0].message.content,
                          model=settings.OPENAI_MODEL,
                          usage=Usage(
                                prompt_tokens = response.usage.prompt_tokens,
                                completion_tokens = response.usage.completion_tokens,
                                total_tokens = response.usage.total_tokens
                          )if response.usage else None
                        )
    return result

@app.get("/sessions/{session_id}/messages")
async def get_message(session_id: int, db: Session = Depends(get_session)):
    if session_id == None:
        return {"session_id": session_id, "error" : "当前session_id不存在"}
    sql = (
        select(ChatMessage)
        .where(ChatMessage.session_id == session_id)
        .order_by(ChatMessage.created_at)
    )
    messages = db.exec(sql).all()
    return {"session_id": session_id, "messages" : messages}
    
@app.post("/documents/upload", response_model=DocumentUploadResponse)
async def upload_document(
    user_id: int,
    file: UploadFile,
    db: Session = Depends(get_session)):

    # todo 优化pdf解析速度

    status = "uploaded"
    start = time.perf_counter()
    # 1. 校验文件名
    if not file.filename:
        raise HTTPException(status_code=400, detail="文件名不能为空")

    # 2. 校验格式
    if not file.filename.lower().endswith(ALLOWED_EXTENSIONS):
        raise HTTPException(status_code=400, detail="仅支持 txt/pdf/docx/xlsx 格式")

    # 3. 读取内容并校验大小（用实际读取的字节判断，file.size 不一定可靠）
    read_start = time.perf_counter()
    content = await file.read()
    read_ms = int((time.perf_counter() - read_start) * 1000)
    if len(content) > MAX_SIZE:
        raise HTTPException(status_code=400, detail=f"文件大小超过限制（最大 {MAX_SIZE // 1024 // 1024}MB）")

    # 4. 提取文本内容
    parse_start = time.perf_counter()
    text_content = None
    ext = Path(file.filename).suffix.lower()
    if ext in (".txt", ".pdf", ".docx", ".xlsx"):
        try:
            text_content = extract_text_from_bytes(content, ext)
        except Exception:
            text_content = None  # 提取失败不阻断上传
            status = "failed"
    parse_ms = int((time.perf_counter() - parse_start) * 1000)


    UPLOAD_DIR.mkdir(exist_ok=True) #若路径不存在则创建
    safe_filename = f"{uuid.uuid4().hex}_{Path(file.filename).name}" #避免文件名重复
    file_path = UPLOAD_DIR / safe_filename #避免文件名重复
    
    #复制文件内容
    write_start = time.perf_counter()
    with file_path.open("wb") as buffer:
        buffer.write(content)
    write_ms = int((time.perf_counter() - write_start) * 1000)

    document = Document(
            user_id=user_id,
            filename=file.filename,
            file_path=str(file_path),
            content_type=file.content_type,
            size=len(content),
            text_content=text_content,
            status=status
        )
    #上传数据库并返回id
    db_start = time.perf_counter()
    db.add(document)
    db.commit()
    db.refresh(document)
    db_ms = int((time.perf_counter() - db_start) * 1000)
    total_ms = int((time.perf_counter() - start) * 1000)
    print(f"upload_document read={read_ms}ms parse={parse_ms}ms write={write_ms}ms db={db_ms}ms total={total_ms}ms file={file.filename}")

    return DocumentUploadResponse(
        document_id=document.id,
        filename=document.filename,
        status=document.status
    )
# unset OPENAI_API_KEY OPENAI_MODEL OPENAI_BASE_URL DATABASE_URL API_KEY APP_NAME ENV

#uvicorn app.main:app --reload --port 18001